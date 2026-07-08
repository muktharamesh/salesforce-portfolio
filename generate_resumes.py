import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def style_paragraph(p, space_before=0, space_after=3, line_spacing=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing

def add_bullet_point(doc, bold_prefix, text, italic_suffix=None):
    p = doc.add_paragraph(style='List Bullet')
    style_paragraph(p, space_before=0, space_after=3)
    
    r_bold = p.add_run(bold_prefix)
    r_bold.bold = True
    r_bold.font.name = 'Arial'
    r_bold.font.size = Pt(9.5)
    r_bold.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    r_text = p.add_run(text)
    r_text.font.name = 'Arial'
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    if italic_suffix:
        r_suffix = p.add_run(italic_suffix)
        r_suffix.italic = True
        r_suffix.font.name = 'Arial'
        r_suffix.font.size = Pt(9.5)
        r_suffix.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

def add_section_header(doc, title):
    p = doc.add_paragraph()
    style_paragraph(p, space_before=12, space_after=4)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(title.upper())
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D) # Navy
    
    # Add a thin bottom border to section title using a paragraph bottom border or horizontal rule in docx is complex,
    # so we will rely on clean spacing and bold upper case.

def build_general_resume():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9.5)
    font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, space_after=2)
    r = p.add_run("MUKTHA RAMESH")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, space_after=3)
    r = p.add_run("Salesforce Administrator & Platform Integrations Specialist")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, space_after=12)
    r = p.add_run("Email: muktha.rameshb@redhorse.red  |  GitHub Portfolio: github.com/muktharamesh/salesforce-portfolio")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    
    # Professional Summary
    add_section_header(doc, "Professional Summary")
    p = doc.add_paragraph()
    style_paragraph(p, space_after=10)
    p.add_run(
        "Result-oriented Salesforce Administrator and Integrations Specialist with a proven track record of designing custom "
        "cloud-to-cloud integrations, database schemas, and robust business automations. Expert in translating business "
        "bottlenecks into scalable technical solutions using Salesforce DX, Lightning Flow, Zapier, and custom metadata "
        "architectures. Proven ability to eliminate manual overhead, improve data integrity by 100%, and deliver real-time "
        "business intelligence to leadership."
    )
    
    # Technical Skills
    add_section_header(doc, "Technical Skills")
    p = doc.add_paragraph()
    style_paragraph(p, space_after=3)
    r = p.add_run("Salesforce Platform: ")
    r.bold = True
    p.add_run("Flow Builder (Screen & Record-Triggered Flows), Custom Objects & Fields, Reports & Dashboards, Salesforce DX CLI, User/Role Management.")
    
    p = doc.add_paragraph()
    style_paragraph(p, space_after=10)
    r = p.add_run("Integrations & Tools: ")
    r.bold = True
    p.add_run("Zapier (NLA, Webhooks, API Mappings), Salesforce Tooling API, REST Webhooks, Git Version Control, GitHub.")
    
    # Key Integrations
    add_section_header(doc, "Key Integrations & Automations (Portfolio Highlights)")
    
    # BetterWorld
    p = doc.add_paragraph()
    style_paragraph(p, space_before=4, space_after=2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("BetterWorld to Salesforce Integration (Cloud-to-Cloud Automation)")
    r.bold = True
    r.font.size = Pt(10)
    
    add_bullet_point(doc, "", "Designed a zero-duplication database landing system connecting BetterWorld auction and donation webhooks to Salesforce via Zapier.")
    add_bullet_point(doc, "Eliminated 100% of manual data entry ", "for all external ticket purchases, donations, and auction winners.")
    add_bullet_point(doc, "Saved 15+ hours per week ", "of manual administration time (approx. 780 hours/year).")
    add_bullet_point(doc, "Saved an estimated $19,500+ in annual operational overhead ", "(based on a standard coordinator rate of $25/hr).")
    add_bullet_point(doc, "Reduced data mapping errors to 0% ", "by building custom validation and landing table architecture.")
    
    # Square
    p = doc.add_paragraph()
    style_paragraph(p, space_before=6, space_after=2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("Square Point-of-Sale Event Integration (Collaborative Project)")
    r.bold = True
    r.font.size = Pt(10)
    
    add_bullet_point(doc, "", "Managed the business architecture and schema mapping for point-of-sale transaction sync (co-developed with Stacy McDonald).")
    add_bullet_point(doc, "Unified in-person card transactions ", "and digital CRM databases, eliminating manual reconciliation sheets for fundraising events.")
    add_bullet_point(doc, "Secured 100% accurate tracking of processing fees ", "and transaction IDs (using Processing_Fee__c), improving financial audit readiness.")
    add_bullet_point(doc, "Saved an estimated 8+ hours ", "of accounting reconciliation time per live event.")
    
    # Custom CSV
    p = doc.add_paragraph()
    style_paragraph(p, space_before=6, space_after=2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("Custom CSV Import Wizard & Record Automator (LWC & Apex Utility)")
    r.bold = True
    r.font.size = Pt(10)
    
    add_bullet_point(doc, "", "Developed a custom Lightning Web Component and Apex database staging framework to automate bulk constituent uploads.")
    add_bullet_point(doc, "Enforced Data Integrity: ", "Created client-side CSV validation and an intermediate staging architecture (Item_Donor__c) with automated Contact matching rules, reducing duplicate accounts by 100%.")
    add_bullet_point(doc, "Automated Gift Processing: ", "Triggered automated Opportunity creation (Record Type: In-Kind Gift) with mapped Campaign sources, Close Dates, realized amounts, and charitable fund allocations (Charitable_Fund__c).")
    add_bullet_point(doc, "Automated Error Tracking: ", "Structured automatic transaction status reporting ('Processed' / 'Failed' with detailed error logs) to allow admins to immediately troubleshoot data issues.")

    # Ghost
    p = doc.add_paragraph()
    style_paragraph(p, space_before=6, space_after=2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("Automated Supporter Segmentation & Privacy Compliance (Ghost Sync Integration)")
    r.bold = True
    r.font.size = Pt(10)
    
    add_bullet_point(doc, "", "Architected a record-triggered flow and database tagging engine connecting Salesforce to the Ghost newsletter and email marketing platform.")
    add_bullet_point(doc, "Dynamic Supporter Segmentation: ", "Designed a custom formula field (Ghost_Labels__c) that dynamically generates subscriber interest tags based on contact picklists and opportunity record types, enabling targeted automated email campaigns.")
    add_bullet_point(doc, "Automated Sync: ", "Real-time sync of new donors, purchasers, and event registrants to Ghost via Zapier, saving 3–5 hours/week of manual list management.")
    add_bullet_point(doc, "Ensured 100% Privacy Compliance: ", "Configured a filter rule to programmatically exclude sensitive Mental Health Session records, eliminating the risk of human error violating privacy policies.")
    
    # Day to Day
    add_section_header(doc, "Day-to-Day Salesforce Administration & Platform Maintenance")
    add_bullet_point(doc, "Database Schema & Custom Fields: ", "Created and updated custom fields, lookups, and formula fields (such as QB_Payment_Method_Formula__c and Current_Calendar_Year__c) to align data formats for downstream accounting systems (QuickBooks).")
    add_bullet_point(doc, "Campaign Hierarchy Management: ", "Structured parent-child campaign roll-ups, ensuring aggregate totals (member counts, donation amounts) roll up accurately for executive visibility.")
    add_bullet_point(doc, "Email Templates: ", "Authored and updated HTML/Classic Email Templates linked to automated system flows to standardize and automate transactional communications.")
    add_bullet_point(doc, "Data Quality & Deduplication: ", "Conducted regular system audits, creating custom reports (like the Duplicate Contact & Account Report) to identify, merge, and purge duplicate records, maintaining a high-fidelity CRM database.")
    add_bullet_point(doc, "Executive Dashboarding: ", "Configured a central executive dashboard fed by 18 custom-built reports tracking year-over-year giving circles, payment offsets, and donor recency, reducing reporting prep time from 4 hours/week to instantaneous, real-time access.")
    
    # Certifications
    add_section_header(doc, "Certifications & Professional Development")
    p = doc.add_paragraph(style='List Bullet')
    style_paragraph(p, space_before=0, space_after=3)
    r = p.add_run("Salesforce Certified Administrator")
    r.font.name = 'Arial'
    
    p = doc.add_paragraph(style='List Bullet')
    style_paragraph(p, space_before=0, space_after=3)
    r = p.add_run("Continuous study in Salesforce Lightning Web Components (LWC) and Apex Development.")
    r.font.name = 'Arial'
    
    doc.save("resume.docx")
    print("Saved resume.docx successfully")

def build_rotary_resume():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9.5)
    font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, space_after=2)
    r = p.add_run("MUKTHA RAMESH")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, space_after=3)
    r = p.add_run("Salesforce Product Owner & Business Analyst")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, space_after=12)
    r = p.add_run("Email: muktha.rameshb@redhorse.red  |  GitHub Portfolio: github.com/muktharamesh/salesforce-portfolio")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    
    # Professional Summary
    add_section_header(doc, "Professional Summary")
    p = doc.add_paragraph()
    style_paragraph(p, space_after=10)
    p.add_run(
        "Strategic and hands-on Salesforce Product Owner & Business Analyst with over 3 years of experience managing the product "
        "lifecycle in Agile environments. Expert at eliciting user requirements, writing clear user stories, managing transparent "
        "backlogs in JIRA, and defining scalable declarative solutions (Flow, Validation Rules, Lightning App Builder). Proven history "
        "of maximizing business value for non-profit and membership models (NPSP/Nonprofit Cloud) by optimizing constituent experiences, "
        "driving participant growth, and reducing technical debt."
    )
    
    # Technical Skills
    add_section_header(doc, "Technical Skills & Methodologies")
    p = doc.add_paragraph()
    style_paragraph(p, space_after=3)
    r = p.add_run("Product Ownership & Agile: ")
    r.bold = True
    p.add_run("Agile/Scrum Mindset, JIRA, Confluence, Backlog Prioritization, User Story Writing, Requirement Elicitation, User Acceptance Testing (UAT).")
    
    p = doc.add_paragraph()
    style_paragraph(p, space_after=10)
    r = p.add_run("Salesforce Platform: ")
    r.bold = True
    p.add_run("Salesforce Flow, Nonprofit Cloud (NPSP), Lightning App Builder, Page Layouts, Data Loader, Security Model (Profiles & Permission Sets), OOTB Config.")
    
    # Highlights
    add_section_header(doc, "Product Owner & Salesforce Experience (Portfolio Highlights)")
    
    # BetterWorld
    p = doc.add_paragraph()
    style_paragraph(p, space_before=4, space_after=2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("BetterWorld to Salesforce Integration (Product Owner & Business Lead)")
    r.bold = True
    r.font.size = Pt(10)
    
    add_bullet_point(doc, "Requirement Elicitation & Backlog Management: ", "Gathered donor database stakeholder requirements and translated them into JIRA user stories to build a webhook integration connecting BetterWorld to Salesforce NPSP.")
    add_bullet_point(doc, "Business Value Maximization: ", "Directed the design of a custom transaction landing table (BetterWorld_Transaction__c) which eliminated 100% of manual data entries for ticket sales and donations, saving 15+ hours/week in operational overhead.")
    add_bullet_point(doc, "Technical Debt Reduction: ", "Prioritized 'Out-of-the-Box' matching rules and native Flows over custom Apex/LWC code, reducing future maintenance overhead.")
    add_bullet_point(doc, "", "Oversaw all data management validation, testing execution, and end-user UAT cycles.")
    
    # Square
    p = doc.add_paragraph()
    style_paragraph(p, space_before=6, space_after=2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("Square Point-of-Sale Integration (Agile Product Owner Collaboration)")
    r.bold = True
    r.font.size = Pt(10)
    
    add_bullet_point(doc, "Cross-Functional Product Collaboration: ", "Collaborated directly with a Salesforce developer to align technical solutions with event team requirements; wrote and accepted user stories for point-of-sale terminal syncs.")
    add_bullet_point(doc, "Process Optimization: ", "Defined acceptance criteria for automated contact matching, eliminating manual spreadsheet reconciliations and saving accounting teams 8+ hours of overhead per live fundraising event.")
    add_bullet_point(doc, "Backlog Prioritization: ", "Managed user story prioritization to balance rapid delivery with platform safety, securing accurate audit logs for processing fees (Processing_Fee__c).")
    
    # Custom CSV
    p = doc.add_paragraph()
    style_paragraph(p, space_before=6, space_after=2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("Custom CSV Import Wizard & Record Automator (LWC & Apex Product Lead)")
    r.bold = True
    r.font.size = Pt(10)
    
    add_bullet_point(doc, "Requirements Elicitation & Staging Architecture: ", "Elicited non-profit gift processing requirements to design an interactive CSV upload tool. Wrote user stories for client-side CSV parsing, an intermediate data staging table (Item_Donor__c), and automated Contact-to-Gift creation.")
    add_bullet_point(doc, "Business Process Optimization: ", "Wrote acceptance criteria for automated deduplication and contact matching, resulting in 100% data integrity and eliminating custom manual deduplication tasks.")
    add_bullet_point(doc, "UAT & Verification Planning: ", "Designed and executed user acceptance testing (UAT) scenarios and verified test suites (ItemDonorImporterTest), ensuring all Opportunity creations (Record Type: In-Kind Gift) and Campaign linkages work flawlessly.")

    # Ghost
    p = doc.add_paragraph()
    style_paragraph(p, space_before=6, space_after=2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("Salesforce to Ghost Member Sync (Privacy & Participant Engagement Product Lead)")
    r.bold = True
    r.font.size = Pt(10)
    
    add_bullet_point(doc, "Customer-Centric UX Design: ", "Elicited requirements to automate newsletter onboarding for new participants. Designed the logical flow to push contact details to Ghost instantly via Zapier, increasing immediate marketing engagement.")
    add_bullet_point(doc, "Risk Triage & Compliance: ", "Structured custom formula criteria (Ghost_Labels__c) to programmatically exclude sensitive medical record types (Mental Health Sessions), ensuring 100% HIPAA and privacy compliance.")
    add_bullet_point(doc, "", "Delivered transparent backlog items and hosted stakeholder demos and training sessions.")

    # Professional Experience
    add_section_header(doc, "Professional Experience")
    p = doc.add_paragraph()
    style_paragraph(p, space_before=4, space_after=2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("Product Compliance Specialist (Consumer Products), Amazon — 2013")
    r.bold = True
    r.font.size = Pt(10)
    
    add_bullet_point(doc, "Requirement Analysis & Technical Translation: ", "Evaluated complex international regulatory frameworks (EU/North American) and translated compliance mandates into actionable technical workflows and product specifications.")
    add_bullet_point(doc, "Change Management & Deployment: ", "Executed policy roll-outs across global teams using structured change management methodologies to ensure smooth operational transitions.")
    add_bullet_point(doc, "Process Standardization: ", "Analyzed operational listing procedures to identify bottlenecks; standardized system workflows to optimize efficiency and meet organizational KPI goals.")
    add_bullet_point(doc, "Stakeholder Collaboration & Risk Triage: ", "Facilitated communications between vendors, testing laboratories, and cross-functional teams to resolve compliance escalations and mitigate project risks.")
    
    # Maintenance
    add_section_header(doc, "Declarative Configuration & Maintenance Work")
    add_bullet_point(doc, "Salesforce Flow Automation: ", "Designed, built, and maintained automated Salesforce Flows (e.g. BetterWorld_Auto_Create) to streamline user journeys and eliminate duplicate records.")
    add_bullet_point(doc, "User Interface Customization: ", "Customized Lightning Record Pages, Page Layouts, List Views, and Compact Layouts to optimize UX for both internal CRM users and external community members.")
    add_bullet_point(doc, "Campaign Hierarchy Architecture: ", "Configured parent-child campaign hierarchies to roll up volunteer hours, attendance metrics, and donation values for real-time executive dashboards.")
    add_bullet_point(doc, "Security & Permissions: ", "Managed user access controls, profiles, custom sharing rules, and permission sets to enforce data governance standards.")
    add_bullet_point(doc, "Analytics & Dashboards: ", "Configured the 2026 RedHorse Dashboard and 18 custom reports, giving leadership immediate access to fundraising, membership growth, and data quality metrics.")
    
    # Education
    add_section_header(doc, "Education & Certifications")
    add_bullet_point(doc, "", "Bachelor's Degree (or equivalent experience)")
    add_bullet_point(doc, "", "Salesforce Certified Administrator")
    add_bullet_point(doc, "", "Scaled Agile - SAFe® Product Owner (In Progress)")
    
    doc.save("resume_rotary.docx")
    print("Saved resume_rotary.docx successfully")

if __name__ == "__main__":
    build_general_resume()
    build_rotary_resume()
